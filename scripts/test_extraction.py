#!/usr/bin/env python3
"""
Simple test script to extract content from one Word document
"""

import sys
import subprocess
from pathlib import Path

def install_package(package_name):
    """Install a package using pip"""
    try:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', package_name])
        return True
    except subprocess.CalledProcessError:
        return False

def test_docx_extraction():
    """Test extraction from a .docx file"""
    try:
        from docx import Document
        print("✓ python-docx is available")
        return True
    except ImportError:
        print("Installing python-docx...")
        if install_package('python-docx'):
            try:
                from docx import Document
                print("✓ python-docx installed successfully")
                return True
            except ImportError:
                print("✗ Failed to install python-docx")
                return False
        return False

def test_doc_extraction():
    """Test extraction from .doc files"""
    try:
        import docx2txt
        print("✓ docx2txt is available")
        return True
    except ImportError:
        print("Installing docx2txt...")
        if install_package('docx2txt'):
            try:
                import docx2txt
                print("✓ docx2txt installed successfully")
                return True
            except ImportError:
                print("✗ Failed to install docx2txt")
                return False
        return False

def extract_sample_file():
    """Extract content from the smallest file as a test"""
    
    # Try the smallest files first
    test_files = [
        "Burra in the News/Eric_s Research - pages notes.docx",
        "Burra in the News/Compiler_s Note 2015.doc"
    ]
    
    for file_path in test_files:
        path = Path(file_path)
        if path.exists():
            print(f"\nTesting extraction from: {path.name}")
            
            if path.suffix.lower() == '.docx':
                if test_docx_extraction():
                    try:
                        from docx import Document
                        doc = Document(path)
                        print(f"  File opened successfully!")
                        print(f"  Number of paragraphs: {len(doc.paragraphs)}")
                        
                        # Show first few paragraphs
                        print("  First few paragraphs:")
                        for i, para in enumerate(doc.paragraphs[:3]):
                            if para.text.strip():
                                print(f"    {i+1}: {para.text.strip()[:100]}...")
                        return True
                    except Exception as e:
                        print(f"  Error: {e}")
            
            elif path.suffix.lower() == '.doc':
                if test_doc_extraction():
                    try:
                        import docx2txt
                        text = docx2txt.process(str(path))
                        print(f"  File processed successfully!")
                        print(f"  Extracted text length: {len(text)} characters")
                        
                        # Show first few lines
                        lines = [line.strip() for line in text.split('\n') if line.strip()]
                        print("  First few lines:")
                        for i, line in enumerate(lines[:3]):
                            print(f"    {i+1}: {line[:100]}...")
                        return True
                    except Exception as e:
                        print(f"  Error: {e}")
    
    print("No test files found or extraction failed")
    return False

def main():
    print("Document Extraction Test")
    print("=" * 30)
    
    # Test package availability
    print("Checking package availability...")
    
    # Test extraction
    if extract_sample_file():
        print("\n✓ Extraction test successful!")
        print("Ready to process all documents.")
    else:
        print("\n✗ Extraction test failed.")
        print("Please check package installation.")

if __name__ == "__main__":
    main()