#!/usr/bin/env python3
"""
Document Extractor for Burra in the News Archive
Extracts text content from .doc and .docx files and converts to structured data
"""

import os
import sys
from pathlib import Path
import json
from typing import Dict, List, Any

def install_requirements():
    """Install required packages if not available"""
    import subprocess
    
    packages = [
        'python-docx',
        'docx2txt'
    ]
    
    for package in packages:
        try:
            __import__(package.replace('-', '_'))
            print(f"✓ {package} already installed")
        except ImportError:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])

def extract_docx_content(file_path: Path) -> Dict[str, Any]:
    """Extract content from .docx files using python-docx"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Only add non-empty paragraphs
                paragraphs.append({
                    'text': text,
                    'style': para.style.name if para.style else 'Normal'
                })
        
        # Extract tables if any
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)
        
        return {
            'file_name': file_path.name,
            'file_type': 'docx',
            'paragraphs': paragraphs,
            'tables': tables,
            'paragraph_count': len(paragraphs),
            'table_count': len(tables)
        }
        
    except Exception as e:
        print(f"Error processing {file_path.name}: {str(e)}")
        return None

def extract_doc_content_simple(file_path: Path) -> Dict[str, Any]:
    """Extract content from .doc files using docx2txt (simple text extraction)"""
    try:
        import docx2txt
        
        # Extract text content
        text = docx2txt.process(str(file_path))
        
        # Split into paragraphs (simple approach)
        paragraphs = []
        for line in text.split('\n'):
            line = line.strip()
            if line:  # Only add non-empty lines
                paragraphs.append({
                    'text': line,
                    'style': 'Normal'
                })
        
        return {
            'file_name': file_path.name,
            'file_type': 'doc',
            'paragraphs': paragraphs,
            'tables': [],  # docx2txt doesn't preserve table structure well
            'paragraph_count': len(paragraphs),
            'table_count': 0
        }
        
    except Exception as e:
        print(f"Error processing {file_path.name}: {str(e)}")
        return None

def extract_all_documents(source_folder: str = "Burra in the News") -> List[Dict[str, Any]]:
    """Extract content from all Word documents in the specified folder"""
    
    source_path = Path(source_folder)
    if not source_path.exists():
        print(f"Error: Folder '{source_folder}' not found")
        return []
    
    documents = []
    
    # Get all .doc and .docx files
    doc_files = list(source_path.glob("*.doc")) + list(source_path.glob("*.docx"))
    
    print(f"Found {len(doc_files)} Word documents to process...")
    
    for doc_file in doc_files:
        print(f"\nProcessing: {doc_file.name}")
        
        if doc_file.suffix.lower() == '.docx':
            content = extract_docx_content(doc_file)
        else:  # .doc file
            content = extract_doc_content_simple(doc_file)
        
        if content:
            documents.append(content)
            print(f"  ✓ Extracted {content['paragraph_count']} paragraphs")
        else:
            print(f"  ✗ Failed to extract content")
    
    return documents

def save_extracted_data(documents: List[Dict[str, Any]], output_file: str = "extracted_content.json"):
    """Save extracted content to JSON file"""
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(documents, f, indent=2, ensure_ascii=False)
        print(f"\n✓ Saved extracted content to {output_file}")
        return True
    except Exception as e:
        print(f"Error saving to {output_file}: {str(e)}")
        return False

def main():
    """Main extraction process"""
    print("Burra in the News - Document Extractor")
    print("=" * 40)
    
    # Install requirements
    print("Checking and installing requirements...")
    install_requirements()
    
    # Extract documents
    print("\nExtracting document content...")
    documents = extract_all_documents()
    
    if documents:
        # Save extracted data
        save_extracted_data(documents)
        
        # Print summary
        print(f"\n📊 Extraction Summary:")
        print(f"   Documents processed: {len(documents)}")
        
        total_paragraphs = sum(doc['paragraph_count'] for doc in documents)
        print(f"   Total paragraphs: {total_paragraphs}")
        
        # Show file breakdown
        print(f"\n📄 Files processed:")
        for doc in documents:
            print(f"   • {doc['file_name']}: {doc['paragraph_count']} paragraphs")
    
    else:
        print("No documents were successfully processed.")

if __name__ == "__main__":
    main()